import os.path

import docx
from docx.opc.constants import CONTENT_TYPE as CT
from docx.package import Package


# noinspection PyUnresolvedReferences
class NamedDocument(docx.document.Document):
    def __init__(self, filename: str):
        if not os.path.exists(filename):
            self._create_empty_document(filename)

        document_part = Package.open(filename).main_document_part
        if document_part.content_type != CT.WML_DOCUMENT_MAIN:
            tmpl = "file '%s' is not a Word file, content type is '%s'"
            raise ValueError(tmpl % (filename, document_part.content_type))

        self.stream = document_part.document
        self.filename = filename

        # noinspection PyProtectedMember
        super().__init__(document_part._element, document_part)

    @staticmethod
    def _create_empty_document(filename: str):
        doc = docx.Document()
        doc.save(filename)

    def save(self, filename: str = None):
        if filename is None:
            filename = self.filename
        super().save(filename)


# noinspection PyUnresolvedReferences
class FlushableDocument(NamedDocument):
    def flush(self):
        self.save()


class ClosableDocument(NamedDocument):
    def close(self):
        self.save()


class HandlerDocument(FlushableDocument, ClosableDocument):
    def write(self, msg: str):
        self.add_paragraph(msg)
